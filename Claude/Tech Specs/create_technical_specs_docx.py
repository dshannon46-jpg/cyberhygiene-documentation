#!/usr/bin/env python3
"""
Create professionally formatted Technical Specifications Document
Includes: Title page, Table of Contents, proper heading styles
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def create_title_page(doc):
    """Create professional title page"""
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_before = Pt(100)
    run = title.add_run("CyberHygiene Production Network")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_before = Pt(12)
    run = subtitle.add_run("Technical Specifications Document")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Version and classification
    doc.add_paragraph()  # Spacer
    version = doc.add_paragraph()
    version.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    version.paragraph_format.space_before = Pt(50)
    run = version.add_run("Version 1.2")
    run.font.size = Pt(16)
    run.font.bold = True

    # Classification
    classification = doc.add_paragraph()
    classification.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    classification.paragraph_format.space_before = Pt(20)
    run = classification.add_run("CONTROLLED UNCLASSIFIED INFORMATION (CUI)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)  # Red

    # Organization
    doc.add_paragraph()  # Spacer
    doc.add_paragraph()  # Spacer
    org = doc.add_paragraph()
    org.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    org.paragraph_format.space_before = Pt(100)
    run = org.add_run("Donald E. Shannon LLC\ndba The Contract Coach")
    run.font.size = Pt(14)

    # Domain
    domain = doc.add_paragraph()
    domain.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    domain.paragraph_format.space_before = Pt(12)
    run = domain.add_run("cyberinabox.net")
    run.font.size = Pt(12)
    run.font.italic = True

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.paragraph_format.space_before = Pt(150)
    run = date_para.add_run(f"October 28, 2025")
    run.font.size = Pt(12)

    # Status
    status = doc.add_paragraph()
    status.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    status.paragraph_format.space_before = Pt(12)
    run = status.add_run("DRAFT - 94% Implementation Complete")
    run.font.size = Pt(11)
    run.font.italic = True

    add_page_break(doc)

def create_toc_page(doc):
    """Create table of contents page"""
    heading = doc.add_heading('TABLE OF CONTENTS', level=1)
    heading.paragraph_format.page_break_before = False

    # Add TOC field code (Word will populate this when opened)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Instruction text
    instruction = doc.add_paragraph()
    instruction.paragraph_format.space_before = Pt(12)
    run = instruction.add_run("Note: ")
    run.font.italic = True
    run.font.bold = True
    run = instruction.add_run("In Microsoft Word, right-click this field and select 'Update Field' to populate the table of contents with page numbers.")
    run.font.italic = True
    run.font.size = Pt(10)

    add_page_break(doc)

def main():
    """Create the technical specifications document"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Create title page
    create_title_page(doc)

    # Create table of contents
    create_toc_page(doc)

    # DOCUMENT CONTROL
    doc.add_heading('DOCUMENT CONTROL', level=1)
    doc.add_paragraph("Document Status: DRAFT - Implementation Phase")
    doc.add_paragraph("Security Classification: CONTROLLED UNCLASSIFIED INFORMATION (CUI)")
    doc.add_paragraph("Distribution: Limited to authorized personnel only")
    doc.add_paragraph()

    # Approval table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name / Title'
    hdr_cells[1].text = 'Role'
    hdr_cells[2].text = 'Signature / Date'

    row1_cells = table.rows[1].cells
    row1_cells[0].text = 'Donald E. Shannon\nOwner/Principal'
    row1_cells[1].text = 'System Owner'
    row1_cells[2].text = '_____________________\nDate: _______________'

    row2_cells = table.rows[2].cells
    row2_cells[0].text = 'Donald E. Shannon\nOwner/Principal'
    row2_cells[1].text = 'ISSO'
    row2_cells[2].text = '_____________________\nDate: _______________'

    doc.add_paragraph()
    doc.add_paragraph("Review Schedule: Quarterly or upon significant system changes")
    doc.add_paragraph("Next Review Date: January 31, 2026")

    # Revision History
    doc.add_heading('Document Revision History', level=2)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Version'
    hdr_cells[1].text = 'Date'
    hdr_cells[2].text = 'Author'
    hdr_cells[3].text = 'Description'

    row1_cells = table.rows[1].cells
    row1_cells[0].text = '1.0'
    row1_cells[1].text = '10/26/2025'
    row1_cells[2].text = 'D. Shannon'
    row1_cells[3].text = 'Initial technical specifications'

    row2_cells = table.rows[2].cells
    row2_cells[0].text = '1.1'
    row2_cells[1].text = '10/28/2025'
    row2_cells[2].text = 'D. Shannon'
    row2_cells[3].text = 'RAID 5, LUKS, Samba, ClamAV updates'

    row3_cells = table.rows[3].cells
    row3_cells[0].text = '1.2'
    row3_cells[1].text = '10/28/2025'
    row3_cells[2].text = 'D. Shannon'
    row3_cells[3].text = 'Wazuh SIEM/XDR deployment, automated backups, complete software stack documentation'

    add_page_break(doc)

    # EXECUTIVE SUMMARY
    doc.add_heading('EXECUTIVE SUMMARY', level=1)

    doc.add_heading('Purpose', level=2)
    doc.add_paragraph(
        "This document provides comprehensive technical specifications for The Contract Coach's "
        "production network infrastructure. This system is designed to process, store, and transmit "
        "Controlled Unclassified Information (CUI) and Federal Contract Information (FCI) in compliance "
        "with NIST SP 800-171 Rev 2 requirements and CMMC Level 2 certification criteria."
    )

    doc.add_heading('System Overview', level=2)
    doc.add_paragraph("System Name: CyberHygiene Production Network (cyberinabox.net)")
    doc.add_paragraph("Implementation Status: 94% Complete (as of October 28, 2025)")
    doc.add_paragraph("Target Completion: December 31, 2025")
    doc.add_paragraph("Compliance: NIST 800-171 Rev 2, FIPS 140-2, CMMC Level 2 Ready")

    doc.add_heading('Infrastructure Summary', level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Medium Shading 1 Accent 1'

    components = [
        ('Total Systems', '4 (1 server, 3 workstations)'),
        ('Total CPU Cores', '16 (4 cores per system)'),
        ('Total RAM', '128 GB (32 GB per system - uniform)'),
        ('Total Storage', '~9 TB usable (3.5 TB boot + 5.5 TB RAID 5)'),
        ('Network Speed', '1 Gbps (All systems Gigabit Ethernet)'),
        ('Security Compliance', '100% (OpenSCAP CUI profile verified)'),
        ('Encryption', '100% (LUKS full disk encryption, FIPS)'),
    ]

    for i, (component, value) in enumerate(components):
        row_cells = table.rows[i].cells
        row_cells[0].text = component
        row_cells[1].text = value

    doc.add_heading('Key Capabilities', level=2)
    capabilities = [
        "FIPS 140-2 validated cryptography on all systems",
        "100% OpenSCAP CUI compliance (105/105 checks passed)",
        "Enterprise identity management (FreeIPA/Kerberos/LDAP)",
        "Security Information and Event Management (Wazuh SIEM/XDR)",
        "Automated vulnerability scanning and patching",
        "File Integrity Monitoring (real-time and scheduled)",
        "Automated daily and weekly backup with bare-metal recovery",
        "Encrypted RAID 5 storage (5.5 TB usable capacity)",
        "Out-of-band management (iLO 5 on critical systems)",
        "Network security (pfSense firewall with IDS/IPS capability)",
    ]

    for capability in capabilities:
        p = doc.add_paragraph(capability, style='List Bullet')

    add_page_break(doc)

    # SECTION 1: HARDWARE SPECIFICATIONS
    doc.add_heading('SECTION 1: HARDWARE SPECIFICATIONS', level=1)

    doc.add_heading('1.1 System Inventory', level=2)

    # dc1 Server
    doc.add_heading('Domain Controller (dc1.cyberinabox.net)', level=3)
    doc.add_paragraph("System Information", style='Heading 4')
    doc.add_paragraph("• Hostname: dc1.cyberinabox.net")
    doc.add_paragraph("• IP Address: 192.168.1.10/24")
    doc.add_paragraph("• Role: Domain Controller, FreeIPA Server, Wazuh Manager, File Server")
    doc.add_paragraph("• OS: Rocky Linux 9.6 Server")

    doc.add_paragraph("Hardware Platform", style='Heading 4')
    doc.add_paragraph("• Model: HP ProLiant MicroServer Gen10 Plus")
    doc.add_paragraph("• Form Factor: Tower Server")
    doc.add_paragraph("• Management: iLO 5 (Integrated Lights-Out) expansion card")

    doc.add_paragraph("Processor", style='Heading 4')
    doc.add_paragraph("• CPU Cores: 4")
    doc.add_paragraph("• Architecture: x86_64")
    doc.add_paragraph("• Virtualization: Enabled")

    doc.add_paragraph("Memory", style='Heading 4')
    doc.add_paragraph("• Total RAM: 32 GB")
    doc.add_paragraph("• Configuration: ECC (Error-Correcting Code)")
    doc.add_paragraph("• Type: DDR4")

    doc.add_paragraph("Storage Configuration", style='Heading 4')
    doc.add_paragraph("Boot Drive (2 TB):")
    storage_items = [
        "/boot/efi - 952 MB (EFI System Partition)",
        "/boot - 7.4 GB",
        "/ (root) - 90 GB (LVM on LUKS encrypted)",
        "/tmp - 15 GB (LVM on LUKS encrypted)",
        "/var - 30 GB (LVM on LUKS encrypted)",
        "/var/log - 15 GB (LVM on LUKS encrypted)",
        "/var/log/audit - 15 GB (LVM on LUKS encrypted)",
        "/home - 239 GB (LVM on LUKS encrypted)",
        "/backup - 931 GB (LVM on LUKS encrypted, daily backups)",
        "/data - 350 GB (LVM on LUKS encrypted, application data)",
        "Swap - 29 GB (encrypted)",
    ]
    for item in storage_items:
        doc.add_paragraph(f"• {item}")

    doc.add_paragraph()
    doc.add_paragraph("RAID Array (3 × 3 TB SATA HDDs):")
    raid_items = [
        "Configuration: RAID 5 (striping with distributed parity)",
        "Usable Capacity: ~5.5 TB (after RAID 5 parity)",
        "Device: /dev/mapper/samba_data",
        "Mount Point: /srv/samba",
        "Encryption: LUKS (FIPS 140-2 compliant)",
        "Purpose: CUI data storage, Samba file sharing, weekly backup storage (ReaR ISOs)",
        "Fault Tolerance: Single drive failure",
        "Rebuild Time: 8-12 hours (estimated)",
    ]
    for item in raid_items:
        doc.add_paragraph(f"• {item}")

    # LabRat Workstation
    doc.add_heading('LabRat Workstation (labrat.cyberinabox.net)', level=3)
    doc.add_paragraph("• Hostname: labrat.cyberinabox.net")
    doc.add_paragraph("• IP Address: 192.168.1.115/24")
    doc.add_paragraph("• Role: Engineering/Testing Workstation")
    doc.add_paragraph("• OS: Rocky Linux 9.6 Workstation")
    doc.add_paragraph("• Model: HP ProLiant MicroServer Gen10 Plus")
    doc.add_paragraph("• CPU: 4 cores, x86_64")
    doc.add_paragraph("• RAM: 32 GB ECC DDR4")
    doc.add_paragraph("• Boot Drive: 512 GB SSD (LUKS encrypted)")
    doc.add_paragraph("• Management: iLO 5 (out-of-band)")

    # Engineering Workstation
    doc.add_heading('Engineering Workstation (engineering.cyberinabox.net)', level=3)
    doc.add_paragraph("• Hostname: engineering.cyberinabox.net")
    doc.add_paragraph("• IP Address: 192.168.1.104/24")
    doc.add_paragraph("• Role: Engineering/CAD Workstation")
    doc.add_paragraph("• OS: Rocky Linux 9.6 Workstation")
    doc.add_paragraph("• Security Status: ✓ Fully hardened (OpenSCAP 100% CUI compliant)")
    doc.add_paragraph("• Model: HP EliteDesk Microcomputer (USFF)")
    doc.add_paragraph("• CPU: Intel Core i5, 4 cores")
    doc.add_paragraph("• RAM: 32 GB DDR4")
    doc.add_paragraph("• Boot Drive: 256 GB SSD (LUKS encrypted)")
    doc.add_paragraph("• Management: Intel AMT/vPro")

    # Accounting Workstation
    doc.add_heading('Accounting Workstation (accounting.cyberinabox.net)', level=3)
    doc.add_paragraph("• Hostname: accounting.cyberinabox.net")
    doc.add_paragraph("• IP Address: 192.168.1.113/24")
    doc.add_paragraph("• Role: Accounting/Financial Workstation")
    doc.add_paragraph("• OS: Rocky Linux 9.6 Workstation")
    doc.add_paragraph("• Security Status: ✓ Fully hardened (OpenSCAP 100% CUI compliant)")
    doc.add_paragraph("• Model: HP EliteDesk Microcomputer (USFF)")
    doc.add_paragraph("• CPU: Intel Core i5, 4 cores")
    doc.add_paragraph("• RAM: 32 GB DDR4")
    doc.add_paragraph("• Boot Drive: 256 GB SSD (LUKS encrypted)")
    doc.add_paragraph("• Management: Intel AMT/vPro")

    add_page_break(doc)

    # SECTION 2: SOFTWARE STACK
    doc.add_heading('SECTION 2: SOFTWARE STACK', level=1)

    doc.add_heading('2.1 Operating System', level=2)
    doc.add_paragraph("Distribution: Rocky Linux (RHEL binary compatible)")
    doc.add_paragraph("Version: 9.6 (Blue Onyx - released October 2025)")
    doc.add_paragraph("Kernel: Linux 5.14.0-570.x.x.el9_6.x86_64")
    doc.add_paragraph("Architecture: x86_64")
    doc.add_paragraph()
    doc.add_paragraph("Security Features:")
    sec_features = [
        "FIPS 140-2 mode enabled (cryptographic validation)",
        "SELinux enforcing mode",
        "Kernel hardening (sysctl configurations)",
        "Secure Boot enabled (UEFI)",
        "TPM 2.0 support",
    ]
    for feature in sec_features:
        doc.add_paragraph(f"• {feature}")

    doc.add_heading('2.2 Identity and Access Management - FreeIPA 4.11.x', level=2)
    doc.add_paragraph("Purpose: Enterprise identity management, authentication, and authorization")
    doc.add_paragraph()
    doc.add_paragraph("Components:")
    freeipa_components = [
        "389 Directory Server (LDAP)",
        "MIT Kerberos (KDC and authentication)",
        "Dogtag Certificate System (PKI/CA)",
        "BIND DNS with DNSSEC",
        "NTP time synchronization",
        "SSSD (System Security Services Daemon)",
    ]
    for component in freeipa_components:
        doc.add_paragraph(f"• {component}")

    doc.add_paragraph()
    doc.add_paragraph("Deployment:")
    doc.add_paragraph("• Server: dc1.cyberinabox.net (192.168.1.10)")
    doc.add_paragraph("• Realm: CYBERINABOX.NET")
    doc.add_paragraph("• Domain: cyberinabox.net")
    doc.add_paragraph("• CA Subject: CN=Certificate Authority,O=CYBERINABOX.NET")

    doc.add_paragraph()
    doc.add_paragraph("Capabilities:")
    freeipa_caps = [
        "Centralized user and group management",
        "Kerberos single sign-on (SSO)",
        "Host-based access control (HBAC)",
        "Sudo rule management",
        "SSH key distribution",
        "Internal certificate authority",
        "DNS management with DNSSEC",
        "Password policies (complexity, expiration, history)",
        "Two-factor authentication ready (OTP support)",
    ]
    for cap in freeipa_caps:
        doc.add_paragraph(f"• {cap}")

    doc.add_heading('2.3 Security Information and Event Management (SIEM)', level=2)
    doc.add_paragraph("Product: Wazuh v4.9.2")
    doc.add_paragraph("Deployment Date: October 28, 2025")
    doc.add_paragraph("Architecture: Manager + Indexer (Dashboard skipped for FIPS compliance)")
    doc.add_paragraph()

    doc.add_paragraph("Components Deployed:", style='Heading 4')
    doc.add_paragraph()
    doc.add_paragraph("Wazuh Manager (dc1.cyberinabox.net):")
    wazuh_mgr = [
        "Service: wazuh-manager",
        "Ports: 1514 (agent), 1515 (cluster), 55000 (API)",
        "Purpose: Log collection, analysis, correlation, alerting",
        "Log Location: /var/ossec/logs/",
        "Configuration: /var/ossec/etc/ossec.conf",
    ]
    for item in wazuh_mgr:
        doc.add_paragraph(f"• {item}")

    doc.add_paragraph()
    doc.add_paragraph("Wazuh Indexer (dc1.cyberinabox.net):")
    wazuh_idx = [
        "Service: wazuh-indexer",
        "Ports: 9200 (HTTP), 9300 (cluster)",
        "Purpose: Alert storage, search, and retrieval",
        "Technology: OpenSearch (Elasticsearch fork)",
        "Storage: /var/lib/wazuh-indexer/",
        "Retention: 90 days (default, configurable)",
    ]
    for item in wazuh_idx:
        doc.add_paragraph(f"• {item}")

    doc.add_paragraph()
    doc.add_paragraph("Active Modules:", style='Heading 4')
    wazuh_modules = [
        "Vulnerability Detection (hourly CVE database updates, package scanning, CVSS scoring)",
        "File Integrity Monitoring (real-time monitoring of /etc, /usr/bin, /boot; SHA256 checksums)",
        "Security Configuration Assessment (CIS Rocky Linux 9 Benchmark compliance)",
        "Log Collection and Analysis (journald, syslog, SSH, sudo, PAM events)",
    ]
    for i, module in enumerate(wazuh_modules, 1):
        doc.add_paragraph(f"{i}. {module}")

    doc.add_paragraph()
    doc.add_paragraph("Alert Configuration:")
    doc.add_paragraph("• Log Level: 3+ (all significant events logged)")
    doc.add_paragraph("• Email Alerts: Disabled (will enable after email server deployment)")
    doc.add_paragraph("• Email Threshold: Level 12+ (critical alerts)")
    doc.add_paragraph("• Alert Formats: Human-readable (alerts.log) and JSON (alerts.json)")
    doc.add_paragraph("• Alert Location: /var/ossec/logs/alerts/")
    doc.add_paragraph()
    doc.add_paragraph("Compliance Mappings (Automatic): NIST SP 800-53, NIST SP 800-171, PCI DSS, HIPAA, GDPR, MITRE ATT&CK")

    doc.add_heading('2.4 Backup and Disaster Recovery', level=2)
    doc.add_paragraph("Product: ReaR (Relax-and-Recover) 2.7")
    doc.add_paragraph("Deployment: Installed on dc1.cyberinabox.net")
    doc.add_paragraph()
    doc.add_paragraph("Capabilities:")
    rear_caps = [
        "Bootable ISO creation (~890 MB)",
        "Full system backup (tar.gz archives)",
        "Bare-metal recovery",
        "LUKS encryption support",
        "LVM layout recreation",
        "Automated scheduling via cron",
    ]
    for cap in rear_caps:
        doc.add_paragraph(f"• {cap}")

    doc.add_paragraph()
    doc.add_paragraph("Backup Schedule:", style='Heading 4')
    doc.add_paragraph()
    doc.add_paragraph("Weekly Full System Backup:")
    doc.add_paragraph("• Schedule: Sunday 3:00 AM")
    doc.add_paragraph("• Target: /srv/samba/backups/ (RAID 5 array)")
    doc.add_paragraph("• Retention: 4 weeks")
    doc.add_paragraph("• Format: Bootable ISO + full backup tar.gz")
    doc.add_paragraph()
    doc.add_paragraph("Daily Critical Files Backup:")
    doc.add_paragraph("• Schedule: 2:00 AM daily")
    doc.add_paragraph("• Target: /backup (931 GB partition on dc1)")
    doc.add_paragraph("• Retention: 30 days")
    doc.add_paragraph("• Contents: FreeIPA config/database, SSL certificates, LUKS keys, system configs, Wazuh configs, user data snapshots")

    doc.add_heading('2.5 Additional Software Components', level=2)

    doc.add_paragraph("Samba 4.19.x (File Sharing)", style='Heading 4')
    doc.add_paragraph("• Share: /srv/samba (5.5 TB RAID 5 array)")
    doc.add_paragraph("• Encryption: SMB3 encryption enforced")
    doc.add_paragraph("• Authentication: FreeIPA/Kerberos integration")
    doc.add_paragraph("• Security: LUKS encryption at rest, audit logging enabled")
    doc.add_paragraph()

    doc.add_paragraph("ClamAV 1.0.x (Antivirus)", style='Heading 4')
    doc.add_paragraph("• Signatures: 27,673+ virus definitions")
    doc.add_paragraph("• Updates: Hourly checks for new definitions")
    doc.add_paragraph("• Integration: Wazuh monitoring for virus detection alerts")
    doc.add_paragraph("• Capabilities: On-demand scanning, email scanning (future), archive extraction")
    doc.add_paragraph()

    doc.add_paragraph("OpenSCAP 1.3.x (Compliance Scanning)", style='Heading 4')
    doc.add_paragraph("• Deployment: All systems (dc1, LabRat, Engineering, Accounting)")
    doc.add_paragraph("• Profile: CUI (Controlled Unclassified Information)")
    doc.add_paragraph("• Compliance Status: 100% on all systems (105/105 checks passed)")
    doc.add_paragraph("• Schedule: Monthly automated scans")
    doc.add_paragraph()

    doc.add_paragraph("LUKS 2.x (Full Disk Encryption)", style='Heading 4')
    doc.add_paragraph("• Algorithm: AES-256-XTS")
    doc.add_paragraph("• FIPS Mode: Enabled (FIPS 140-2 validated)")
    doc.add_paragraph("• Deployment: All boot drives and RAID 5 array encrypted")
    doc.add_paragraph("• Key Management: TPM 2.0 storage, emergency backups")

    add_page_break(doc)

    # SECTION 3: NETWORK ARCHITECTURE
    doc.add_heading('SECTION 3: NETWORK ARCHITECTURE', level=1)

    doc.add_heading('3.1 Network Topology', level=2)
    doc.add_paragraph("Network Segment: 192.168.1.0/24")
    doc.add_paragraph("Gateway: 192.168.1.1 (pfSense firewall)")
    doc.add_paragraph("Subnet Mask: 255.255.255.0")
    doc.add_paragraph("DNS Server: 192.168.1.10 (dc1 - FreeIPA BIND)")
    doc.add_paragraph("NTP Server: 192.168.1.10 (dc1 - Chrony)")
    doc.add_paragraph()

    doc.add_paragraph("IP Address Assignments:")
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'System'
    hdr_cells[1].text = 'Hostname'
    hdr_cells[2].text = 'IP Address'

    ip_assignments = [
        ('Firewall', 'pfSense', '192.168.1.1'),
        ('Domain Controller', 'dc1.cyberinabox.net', '192.168.1.10'),
        ('Engineering WS', 'engineering.cyberinabox.net', '192.168.1.104'),
        ('Accounting WS', 'accounting.cyberinabox.net', '192.168.1.113'),
        ('LabRat WS', 'labrat.cyberinabox.net', '192.168.1.115'),
    ]

    for i, (system, hostname, ip) in enumerate(ip_assignments, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = system
        row_cells[1].text = hostname
        row_cells[2].text = ip

    doc.add_heading('3.2 Network Security', level=2)
    doc.add_paragraph("Firewall/Router: Netgate 2100 pfSense Appliance")
    doc.add_paragraph("• WAN: 96.72.6.225 (public IP)")
    doc.add_paragraph("• LAN: 192.168.1.1/24")
    doc.add_paragraph("• Features: Firewall, VPN, IDS/IPS (Suricata capable), NAT, DHCP")
    doc.add_paragraph()
    doc.add_paragraph("Encryption in Transit (All network traffic encrypted):")
    encryption = [
        "HTTPS (TLS 1.2+) for web services",
        "SSH (OpenSSH 8.7+) for remote administration",
        "LDAPS (LDAP over TLS) for directory queries",
        "Kerberos (encrypted tickets)",
        "SMB3 encryption for file sharing",
        "DNSSEC for DNS integrity",
    ]
    for item in encryption:
        doc.add_paragraph(f"• {item}")

    add_page_break(doc)

    # SECTION 4: SECURITY ARCHITECTURE
    doc.add_heading('SECTION 4: SECURITY ARCHITECTURE', level=1)

    doc.add_heading('4.1 Defense in Depth (7 Layers)', level=2)
    layers = [
        ("Layer 1: Physical Security", ["Locked server room", "Chassis intrusion detection", "Cable locks", "Facility alarm system"]),
        ("Layer 2: Network Security", ["pfSense firewall", "IDS/IPS capability (Suricata)", "Network segmentation", "Firewall logging"]),
        ("Layer 3: Host Security", ["FIPS 140-2 cryptography", "SELinux enforcing", "LUKS full disk encryption", "Secure Boot", "TPM 2.0"]),
        ("Layer 4: Application Security", ["TLS encryption", "Kerberos SSO", "Role-based access control", "Least privilege"]),
        ("Layer 5: Data Security", ["Encryption at rest (LUKS)", "Encryption in transit (TLS/SSH/SMB3)", "CUI data markings", "Backup and recovery"]),
        ("Layer 6: Monitoring and Detection", ["Wazuh SIEM/XDR", "File Integrity Monitoring", "Vulnerability scanning", "Audit logging"]),
        ("Layer 7: Incident Response", ["Automated alerting", "Active response", "Backup/recovery procedures", "Forensic logging"]),
    ]

    for layer_name, items in layers:
        doc.add_paragraph(layer_name, style='Heading 4')
        for item in items:
            doc.add_paragraph(f"• {item}")

    doc.add_heading('4.2 Access Control', level=2)
    doc.add_paragraph("Authentication Methods:")
    doc.add_paragraph("• Primary: Kerberos SSO (FreeIPA)")
    doc.add_paragraph("• Password Policy: 14+ characters, complexity requirements, 90-day expiration")
    doc.add_paragraph("• Account Lockout: 5 failed attempts, 30-minute lockout")
    doc.add_paragraph("• Administrative: SSH key-based authentication (password disabled for root)")
    doc.add_paragraph("• Future Enhancement: Multi-factor authentication (OTP, PIV/CAC)")
    doc.add_paragraph()
    doc.add_paragraph("Authorization:")
    doc.add_paragraph("• Role-Based Access Control (RBAC) via FreeIPA groups")
    doc.add_paragraph("• Sudo rules by group")
    doc.add_paragraph("• Host-based access control (HBAC)")
    doc.add_paragraph("• SELinux role-based separation")
    doc.add_paragraph("• Principle of least privilege enforced")

    add_page_break(doc)

    # SECTION 5: COMPLIANCE STATUS
    doc.add_heading('SECTION 5: COMPLIANCE STATUS', level=1)

    doc.add_heading('5.1 NIST SP 800-171 Rev 2 Compliance', level=2)
    doc.add_paragraph("Implementation Summary:")
    doc.add_paragraph("• Total Controls: 110 (14 families)")
    doc.add_paragraph("• Implemented: 103 controls (94%)")
    doc.add_paragraph("• Partially Implemented: 7 controls (6%)")
    doc.add_paragraph("• Not Implemented: 0 controls")
    doc.add_paragraph("• Estimated SPRS Score: ~91 points (out of 110)")
    doc.add_paragraph()

    doc.add_paragraph("Control Family Implementation Status:")
    table = doc.add_table(rows=15, cols=3)
    table.style = 'Medium Shading 1 Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Family'
    hdr_cells[1].text = 'Controls'
    hdr_cells[2].text = 'Status'

    families = [
        ('Access Control (AC)', '22', '100% (22/22)'),
        ('Awareness and Training (AT)', '3', '33% (1/3)'),
        ('Audit and Accountability (AU)', '9', '100% (9/9)'),
        ('Configuration Management (CM)', '9', '100% (9/9)'),
        ('Identification and Authentication (IA)', '11', '91% (10/11)'),
        ('Incident Response (IR)', '4', '75% (3/4)'),
        ('Maintenance (MA)', '6', '100% (6/6)'),
        ('Media Protection (MP)', '8', '100% (8/8)'),
        ('Personnel Security (PS)', '2', '100% (2/2)'),
        ('Physical Protection (PE)', '6', '100% (6/6)'),
        ('Risk Assessment (RA)', '3', '100% (3/3)'),
        ('Security Assessment (CA)', '4', '100% (4/4)'),
        ('System and Communications Protection (SC)', '15', '100% (15/15)'),
        ('System and Information Integrity (SI)', '8', '100% (8/8)'),
    ]

    for i, (family, controls, status) in enumerate(families, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = family
        row_cells[1].text = controls
        row_cells[2].text = status

    doc.add_heading('5.2 FIPS 140-2 Compliance', level=2)
    doc.add_paragraph("FIPS Mode Status: ✓ ENABLED on all systems")
    doc.add_paragraph()
    doc.add_paragraph("Cryptographic Modules:")
    doc.add_paragraph("• Kernel Crypto API (AES, SHA, HMAC)")
    doc.add_paragraph("• OpenSSL 3.0 (FIPS module)")
    doc.add_paragraph("• GnuTLS (FIPS mode)")
    doc.add_paragraph("• NSS (Mozilla Network Security Services, FIPS mode)")
    doc.add_paragraph()
    doc.add_paragraph("Algorithms Used:")
    doc.add_paragraph("• Symmetric: AES-256-XTS (LUKS), AES-256-GCM (TLS)")
    doc.add_paragraph("• Hashing: SHA-256, SHA-512")
    doc.add_paragraph("• HMAC: HMAC-SHA-256, HMAC-SHA-512")
    doc.add_paragraph("• Key Exchange: ECDHE, DHE (TLS)")
    doc.add_paragraph("• Digital Signature: RSA-2048, ECDSA")

    doc.add_heading('5.3 OpenSCAP CUI Profile Compliance', level=2)
    doc.add_paragraph("All Systems: 100% Compliant")
    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'System'
    hdr_cells[1].text = 'Profile'
    hdr_cells[2].text = 'Compliance'

    scap_results = [
        ('dc1.cyberinabox.net', 'SCAP - CUI', '105/105 (100%)'),
        ('labrat.cyberinabox.net', 'SCAP - CUI', '105/105 (100%)'),
        ('engineering.cyberinabox.net', 'SCAP - CUI', '105/105 (100%)'),
        ('accounting.cyberinabox.net', 'SCAP - CUI', '105/105 (100%)'),
    ]

    for i, (system, profile, compliance) in enumerate(scap_results, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = system
        row_cells[1].text = profile
        row_cells[2].text = compliance

    doc.add_paragraph()
    doc.add_paragraph("Last Scan: October 28, 2025")

    doc.add_heading('5.4 CMMC Level 2 Readiness', level=2)
    doc.add_paragraph("CMMC 2.0 Level 2 Assessment:")
    doc.add_paragraph("• Total Practices: 110 (same as NIST 800-171)")
    doc.add_paragraph("• Implementation Status: 94% (103/110 practices)")
    doc.add_paragraph("• Maturity Level 1 (Performed): ✓ Met")
    doc.add_paragraph("• Maturity Level 2 (Documented): ✓ Met (SSP/POAM documented)")
    doc.add_paragraph("• Maturity Level 3 (Managed): Planned (continuous improvement)")
    doc.add_paragraph("• CMMC Assessment Readiness: January 1, 2026 (after POA&M completion)")

    add_page_break(doc)

    # APPENDICES
    doc.add_heading('APPENDIX A: Service Port Matrix', level=1)
    table = doc.add_table(rows=13, cols=5)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Service'
    hdr_cells[1].text = 'Port(s)'
    hdr_cells[2].text = 'Protocol'
    hdr_cells[3].text = 'System'
    hdr_cells[4].text = 'Purpose'

    services = [
        ('DNS', '53', 'TCP/UDP', 'dc1', 'Domain name resolution'),
        ('DHCP', '67', 'UDP', 'pfSense', 'IP address assignment'),
        ('Kerberos', '88', 'TCP/UDP', 'dc1', 'Authentication'),
        ('NTP', '123', 'UDP', 'dc1', 'Time synchronization'),
        ('LDAP', '389', 'TCP', 'dc1', 'Directory services'),
        ('HTTPS', '443', 'TCP', 'dc1, pfSense', 'Web services (encrypted)'),
        ('SMB', '445', 'TCP', 'dc1', 'File sharing'),
        ('Kerberos Passwd', '464', 'TCP/UDP', 'dc1', 'Password changes'),
        ('LDAPS', '636', 'TCP', 'dc1', 'Encrypted directory'),
        ('Wazuh Agent', '1514', 'TCP', 'dc1', 'Security monitoring'),
        ('Wazuh Cluster', '1515', 'TCP', 'dc1', 'Cluster communication'),
        ('Wazuh Indexer', '9200', 'TCP', 'dc1', 'Alert storage (internal)'),
        ('Wazuh API', '55000', 'TCP', 'dc1', 'API access'),
    ]

    for i, (service, port, protocol, system, purpose) in enumerate(services, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = service
        row_cells[1].text = port
        row_cells[2].text = protocol
        row_cells[3].text = system
        row_cells[4].text = purpose

    add_page_break(doc)

    doc.add_heading('APPENDIX B: Acronyms and Abbreviations', level=1)
    table = doc.add_table(rows=51, cols=2)
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Acronym'
    hdr_cells[1].text = 'Definition'

    acronyms = [
        ('AC', 'Access Control'),
        ('AES', 'Advanced Encryption Standard'),
        ('AMT', 'Active Management Technology (Intel)'),
        ('AU', 'Audit and Accountability'),
        ('CA', 'Certificate Authority, Security Assessment'),
        ('CAGE', 'Commercial and Government Entity'),
        ('CIS', 'Center for Internet Security'),
        ('CM', 'Configuration Management'),
        ('CMMC', 'Cybersecurity Maturity Model Certification'),
        ('CPU', 'Central Processing Unit'),
        ('CUI', 'Controlled Unclassified Information'),
        ('CVE', 'Common Vulnerabilities and Exposures'),
        ('CVSS', 'Common Vulnerability Scoring System'),
        ('DFARS', 'Defense Federal Acquisition Regulation Supplement'),
        ('DHCP', 'Dynamic Host Configuration Protocol'),
        ('DNS', 'Domain Name System'),
        ('DNSSEC', 'DNS Security Extensions'),
        ('ECC', 'Error-Correcting Code'),
        ('FAR', 'Federal Acquisition Regulation'),
        ('FCI', 'Federal Contract Information'),
        ('FIM', 'File Integrity Monitoring'),
        ('FIPS', 'Federal Information Processing Standards'),
        ('GB', 'Gigabyte'),
        ('GDPR', 'General Data Protection Regulation'),
        ('GNOME', 'GNU Network Object Model Environment'),
        ('HBAC', 'Host-Based Access Control'),
        ('HIPAA', 'Health Insurance Portability and Accountability Act'),
        ('HTTPS', 'Hypertext Transfer Protocol Secure'),
        ('IA', 'Identification and Authentication'),
        ('IDS', 'Intrusion Detection System'),
        ('iLO', 'Integrated Lights-Out (HP management)'),
        ('IP', 'Internet Protocol'),
        ('IPS', 'Intrusion Prevention System'),
        ('IR', 'Incident Response'),
        ('ISSO', 'Information System Security Officer'),
        ('KDC', 'Key Distribution Center (Kerberos)'),
        ('LDAP', 'Lightweight Directory Access Protocol'),
        ('LUKS', 'Linux Unified Key Setup'),
        ('LVM', 'Logical Volume Manager'),
        ('MA', 'Maintenance'),
        ('MFA', 'Multi-Factor Authentication'),
        ('MP', 'Media Protection'),
        ('NIST', 'National Institute of Standards and Technology'),
        ('NTP', 'Network Time Protocol'),
        ('OTP', 'One-Time Password'),
        ('PE', 'Physical Protection'),
        ('PKI', 'Public Key Infrastructure'),
        ('POA&M', 'Plan of Action and Milestones'),
        ('RAID', 'Redundant Array of Independent Disks'),
        ('RAM', 'Random Access Memory'),
    ]

    for i, (acronym, definition) in enumerate(acronyms, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = acronym
        row_cells[1].text = definition

    add_page_break(doc)

    # Final approval page
    doc.add_heading('DOCUMENT APPROVAL', level=1)
    doc.add_paragraph(
        "This Technical Specifications Document accurately represents the CyberHygiene Production "
        "Network as deployed and configured as of October 28, 2025."
    )
    doc.add_paragraph()
    doc.add_paragraph("System Owner:")
    doc.add_paragraph()
    doc.add_paragraph("Signature: _______________________________  Date: _______________")
    doc.add_paragraph("Donald E. Shannon, Owner/Principal")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("ISSO:")
    doc.add_paragraph()
    doc.add_paragraph("Signature: _______________________________  Date: _______________")
    doc.add_paragraph("Donald E. Shannon, Information System Security Officer")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Footer information
    footer_para = doc.add_paragraph()
    footer_para.paragraph_format.space_before = Pt(50)
    run = footer_para.add_run("Last Updated: October 28, 2025")
    run.font.size = Pt(10)
    run.font.italic = True
    doc.add_paragraph("Document Version: 1.2")
    doc.add_paragraph("Classification: CONTROLLED UNCLASSIFIED INFORMATION (CUI)")
    doc.add_paragraph("Distribution: Limited to authorized personnel only")
    doc.add_paragraph()

    end_para = doc.add_paragraph()
    end_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = end_para.add_run("END OF DOCUMENT")
    run.font.bold = True
    run.font.size = Pt(12)

    # Save the document
    output_path = "/home/dshannon/Documents/Claude/CyberHygiene_Technical_Specifications_v1.2_Professional.docx"
    doc.save(output_path)
    print(f"Document created successfully: {output_path}")

if __name__ == "__main__":
    main()
